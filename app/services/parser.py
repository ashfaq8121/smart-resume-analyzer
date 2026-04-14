# parser.py - Extracts raw text from PDF and DOCX resume files
# This is the FIRST step: before we can analyze anything, we need the text.

import re
import io

# For reading PDF files
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Better PDF reader (handles complex layouts)
try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

# For reading Word documents
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResumeParser:
    """
    Handles reading resume files (PDF/DOCX) and returning clean text.
    
    Think of this class like a "universal reader" that can open
    different file formats and give you the text inside.
    """

    def extract_text(self, file_path: str) -> str:
        """
        Main entry point. Given a file path, figure out the type
        and extract text appropriately.
        
        Args:
            file_path: Path to the resume file on disk
            
        Returns:
            Cleaned text string from the resume
        """
        # Check the file extension (lowercase for safety)
        extension = file_path.lower().split('.')[-1]
        
        if extension == 'pdf':
            text = self._extract_from_pdf(file_path)
        elif extension == 'docx':
            text = self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: .{extension}. Use PDF or DOCX.")
        
        # Clean the extracted text before returning
        return self._clean_text(text)

    def extract_text_from_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        Extract text directly from file bytes (useful when file is in memory).
        Used when Flask receives uploaded files without saving first.
        """
        extension = filename.lower().split('.')[-1]
        
        if extension == 'pdf':
            text = self._extract_pdf_from_bytes(file_bytes)
        elif extension == 'docx':
            text = self._extract_docx_from_bytes(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: .{extension}")
        
        return self._clean_text(text)

    # ─────────────────────────────────────────────
    # PDF EXTRACTION
    # ─────────────────────────────────────────────

    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Try pdfminer first (better quality), fall back to PyPDF2.
        Having two methods = more resumes work correctly.
        """
        text = ""
        
        # Method 1: pdfminer (handles complex PDFs better)
        if PDFMINER_AVAILABLE:
            try:
                text = pdfminer_extract(file_path)
                if text and len(text.strip()) > 50:  # Got meaningful text
                    return text
            except Exception:
                pass  # Move to fallback
        
        # Method 2: PyPDF2 (simpler but widely supported)
        if PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:  # 'rb' = read as binary
                    reader = PyPDF2.PdfReader(f)
                    
                    # Loop through every page and collect text
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                raise ValueError(f"Could not read PDF: {e}")
        
        if not text:
            raise ValueError("PDF appears to be empty or image-only (scanned PDF).")
        
        return text

    def _extract_pdf_from_bytes(self, file_bytes: bytes) -> str:
        """Extract PDF text from raw bytes (in-memory file)."""
        text = ""
        
        if PDFMINER_AVAILABLE:
            try:
                text = pdfminer_extract(io.BytesIO(file_bytes))
                if text and len(text.strip()) > 50:
                    return text
            except Exception:
                pass
        
        if PYPDF2_AVAILABLE:
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e:
                raise ValueError(f"Could not read PDF bytes: {e}")
        
        return text

    # ─────────────────────────────────────────────
    # DOCX EXTRACTION
    # ─────────────────────────────────────────────

    def _extract_from_docx(self, file_path: str) -> str:
        """
        Read a Word (.docx) file and extract all paragraph text.
        Also extracts text from tables (some resumes use table layouts).
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from regular paragraphs
            # Each paragraph is a line/block of text in Word
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    text_parts.append(paragraph.text)
            
            # Extract text from tables (some resumes use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        
        except Exception as e:
            raise ValueError(f"Could not read DOCX: {e}")

    def _extract_docx_from_bytes(self, file_bytes: bytes) -> str:
        """Extract DOCX text from raw bytes."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed.")
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Could not read DOCX bytes: {e}")

    # ─────────────────────────────────────────────
    # TEXT CLEANING
    # ─────────────────────────────────────────────

    def _clean_text(self, text: str) -> str:
        """
        Clean up extracted text:
        - Remove weird characters from PDF parsing
        - Normalize whitespace (multiple spaces → single space)
        - Remove non-printable characters
        
        This is important because raw PDF text is often messy!
        """
        if not text:
            return ""
        
        # Replace common PDF artifacts (weird hyphens, line breaks)
        text = text.replace('\x00', ' ')    # Null bytes
        text = text.replace('\uf0b7', ' ')  # Bullet point artifacts
        text = text.replace('\r\n', '\n')   # Windows line endings
        text = text.replace('\r', '\n')     # Old Mac line endings
        
        # Remove non-printable characters but keep newlines and tabs
        # \x20-\x7E is the range of printable ASCII characters
        text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)
        
        # Normalize multiple spaces to single space
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Normalize multiple newlines to maximum 2
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove trailing/leading whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)  # Remove empty lines
        
        return text.strip()
